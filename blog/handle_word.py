# -*- encoding: utf-8 -*-
import csv
import datetime
import os
import re
import subprocess
from pprint import pprint

import docx
import pandas as pd
from docx.shared import Pt
from win32com import client as wc

# Target Name列中如果出现软件名重复列，默认保留遇到的第一个配置,如下为excel第一行必有标题列字段
COLUMNS_NAME, COLUMNS_LICENSE, CONDITION = ['Target Name', 'Target License', '5G是否使用']


def read_excel(excel_path):
    df = pd.read_excel(excel_path)
    if COLUMNS_NAME not in df.columns or COLUMNS_LICENSE not in df.columns:
        raise Exception('请确认excel第一行的软件名配置列包含{}、{}'.format(COLUMNS_NAME, COLUMNS_LICENSE))
    if CONDITION in df.columns:
        target = df[df[CONDITION].isin(['是', 'Y'])][[COLUMNS_NAME, COLUMNS_LICENSE]]
    else:
        target = df[[COLUMNS_NAME, COLUMNS_LICENSE]]
    # target[COLUMNS_NAME] = target[COLUMNS_NAME].str.lower()
    target = target.drop_duplicates(subset=[COLUMNS_NAME])

    soft_name = [i.strip() for i in target[COLUMNS_NAME]]
    soft_license = [re.sub(r'[\n]', ' ', i).strip() for i in target[COLUMNS_LICENSE].fillna('')]
    license_map = dict(zip(soft_name, soft_license))
    return license_map


# 先转换成docx，模块只能打开docx
def read_word(word_path):
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(word_path)

    dir_name, file_name = os.path.split(word_path)
    out_path = os.path.join(os.path.dirname(dir_name), 'output')
    if not os.path.exists(out_path):
        os.makedirs(out_path)

    new_file_path = os.path.join(out_path, file_name)
    doc.SaveAs(new_file_path, 16)
    doc.Close()
    word.Quit()

    doc = docx.Document(new_file_path)
    return doc, new_file_path


def out_csv(path, inconsistent, redundant_software, lack_of_soft):
    file_name = os.path.basename(path).split('.')[0]
    path = os.path.join(os.path.dirname(os.path.abspath(path)), file_name + '_difference.csv')
    with open(path, 'w', encoding='utf_8_sig', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',', quoting=csv.QUOTE_ALL)
        flag = True
        writer.writerow(['license不一致的software', ])
        for row in inconsistent:
            if flag:
                keys = list(row.keys())
                writer.writerow(keys)
                flag = False
            writer.writerow(list(row.values()))

        writer.writerow(['文档中冗余的software', '\n'])
        writer.writerow(redundant_software)
        writer.writerow(['文档中缺少的software', '\n'])
        writer.writerow(lack_of_soft)
        writer.writerow([datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S %f'), ])
        csvfile.close()


def compare_and_fix_word(doc, license_map, new_path):
    soft_obj = re.compile(r'Software[：:]\s*(.*)')
    license_obj = re.compile(r'License[：:]\s*(.*)')

    current_soft = ''
    # 匹配结果放入3个容器中
    consistent = []
    inconsistent = []
    redundant_software = []
    all_software = []

    style = doc.styles['Default']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)

    # 提取word内容, 这里是以段落为单位的，下面用for遍历每一个段落
    for index, p in enumerate(doc.paragraphs):
        ret = soft_obj.match(p.text.strip())
        if ret:
            all_software.append(ret.group(1))
            current_soft = ret.group(1)
        ret = license_obj.match(p.text)
        if ret:
            # word中冗余的软件
            config_license = license_map.get(current_soft, '')
            if not config_license:
                redundant_software.append(current_soft)
                continue

            cu_license = ret.group(1).strip()


            # license匹配
            if cu_license == config_license:
                consistent.append(current_soft)
                doc.paragraphs[index].text = ''
                doc.paragraphs[index].add_run('License: ').bold = True
                doc.paragraphs[index].add_run(cu_license)
            # license不匹配，记录并修改
            else:
                inconsistent.append({"soft_name": current_soft, "old_license": cu_license, 'new_license': config_license})
                doc.paragraphs[index].text = ''
                doc.paragraphs[index].add_run('License: ').bold = True
                doc.paragraphs[index].add_run(config_license)

    doc.save(new_path)

    # 输出操作


    print('*' * 40, 'license一致的共计:', len(consistent))
    print('\t\n'.join(consistent))
    print('*' * 40, 'license不一致的共计:', len(inconsistent))
    pprint(inconsistent)
    print('*' * 40, '文档中多余的软件共计:', len(redundant_software))
    print('\t\n'.join(redundant_software))
    lack_of_soft = sorted(list(set(license_map.keys()) - set([i for i in all_software])))
    print('*' * 40, '文档中缺少的软件共计:', len(lack_of_soft))
    print('\t\n'.join(lack_of_soft))



    out_csv(new_path, inconsistent, redundant_software, lack_of_soft)


def prepare_source():
    excel_url = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'excel_source')
    if not os.path.exists(excel_url):
        os.makedirs(excel_url)

    word_url = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'word_source')
    if not os.path.exists(word_url):
        os.makedirs(word_url)

    global EXCEL_PATH, WORD_PATH_LIST
    excel_list = [i for i in os.listdir(excel_url) if i.endswith('.xlsx')]
    if not excel_list:
        print('请放入一个.xlsx结尾的配置文件')
        exit(1)
    elif len(excel_list) > 1:
        print('请放入唯一一个.xlsx结尾的配置文件,或者关闭已经打开的excel文件')
        exit(1)
    else:
        EXCEL_PATH = os.path.join(excel_url, excel_list[0])

    if not os.listdir(word_url):
        print('请放入.doc或docx结尾的核对文档')
        exit(1)
    else:
        WORD_PATH_LIST = [os.path.join(word_url, i) for i in os.listdir(word_url) if i.endswith(('.doc', '.docx'))]


if __name__ == '__main__':

    prepare_source()

    print('read {}...'.format(EXCEL_PATH))
    license_map = read_excel(EXCEL_PATH)

    for WORD_PATH in WORD_PATH_LIST:
        print('read {}...'.format(WORD_PATH))
        doc, outpath = read_word(WORD_PATH)
        print('compare...')
        compare_and_fix_word(doc, license_map, outpath)
        subprocess.call("pause", shell=True)
